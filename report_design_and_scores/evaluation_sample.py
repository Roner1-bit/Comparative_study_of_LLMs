import os
import docx
from docx import Document
from docx.shared import Inches
from PIL import Image

def read_docx(file_path):
    doc = docx.Document(file_path)
    return "\n".join([para.text for para in doc.paragraphs])

cases_dir = r'D:\final_drive_2\(final)merged_images_with_labels_order_and_folders_mask_normalized-20250218T095105Z-001\(final)merged_images_with_labels_order_and_folders_mask_normalized'
plots_dir = r'D:\final_drive_2\(final)merged_images_with_labels_order_and_folders_mask_normalized-20250218T095105Z-001\plots'
output_dir = r'D:\final_drive_2\final_combined_o1_r1_docx'

os.makedirs(output_dir, exist_ok=True)

# Example model encodings
model_encodings = {
    # "deepseek-vl-small.txt": "Model A",
    # "deepseek-vl-tiny.txt": "Model B",
    "deepseek-vl2-api.txt": "Model C",
    "llama-11B.txt": "Model E",
    # "llava-med-response11.txt": "Model F",
    # "nvlm-72b-response.txt": "Model H",
    # "paligemma2-10b-report.txt": "Model I",
    # "paligemma2-28b-report.txt": "Model J",
    # "paligemma2-3b-report.txt": "Model K",
    # "phi-3.5-vision-instruct-response.txt": "Model L",
    # "qwen-vl-2b-response.txt": "Model M",
    "qwen-vl-72b-response.txt": "Model N",
    # "qwen-vl-7b-response.txt": "Model O",
    "meta-vision-90b.txt": "Model Z"
}

for case in os.listdir(cases_dir):
    case_dir = os.path.join(cases_dir, case)
    if not os.path.isdir(case_dir):
        continue

    reference_path = os.path.join(case_dir, f"{case}.docx")
    if not os.path.exists(reference_path):
        continue
    reference_content = read_docx(reference_path)

    for model_file, _ in model_encodings.items():
        model_file_path = os.path.join(case_dir, model_file)
        if not os.path.exists(model_file_path):
            continue

        with open(model_file_path, 'r', encoding='utf-8') as f_model:
            ai_generated = f_model.read()

        file_prefix = model_file.split('.')[0]
        plot_subdir = os.path.join(plots_dir, file_prefix, case)
        o1_eval_path = os.path.join(plot_subdir, f"{case}_{file_prefix}_o1-score.txt")
        r1_eval_path = os.path.join(plot_subdir, f"{case}_{file_prefix}_r1-score.txt")

        o1_content = ""
        if os.path.exists(o1_eval_path):
            with open(o1_eval_path, 'r', encoding='utf-8') as f_o1:
                o1_content = f_o1.read()

        r1_content = ""
        if os.path.exists(r1_eval_path):
            with open(r1_eval_path, 'r', encoding='utf-8') as f_r1:
                r1_content = f_r1.read()

        if not o1_content and not r1_content:
            continue

        # Create Word document
        doc = Document()

        # Insert any PNG images from the patient's folder, rotating them 90° clockwise
        if os.path.isdir(case_dir):
            for img_file in os.listdir(case_dir):
                if img_file.lower().endswith(".png"):
                    img_path = os.path.join(case_dir, img_file)
                    with Image.open(img_path) as img:
                        rotated_img = img.rotate(-90, expand=True)
                        rotated_path = os.path.join(case_dir, f"rotated_{img_file}")
                        rotated_img.save(rotated_path)
                    doc.add_picture(rotated_path, width=Inches(5.0))
                    os.remove(rotated_path)

        doc.add_heading("Reference Report", level=1)
        reference_lines = reference_content.split("\n")
        for line in reference_lines:
            # Example logic: treat lines ending with a colon as special titles
            if line.strip().endswith(":"):
                doc.add_heading("### " + line.strip(), level=1)
            else:
                doc.add_paragraph(line.strip())

        doc.add_heading("AI-Generated Report", level=1)
        doc.add_paragraph(ai_generated)

        doc.add_heading("O1 Evaluation", level=1)
        doc.add_paragraph(o1_content)

        doc.add_heading("R1 Evaluation", level=1)
        doc.add_paragraph(r1_content)

        doc.add_heading("Radiologist Questions", level=1)
        doc.add_paragraph("1. Which evaluator is more accurate in terms of Correctness?\n\n")
        doc.add_paragraph("2. Which evaluator is more accurate in terms of Conciseness?\n\n")
        doc.add_paragraph("3. Which evaluator is more accurate in terms of Completeness?\n\n")
        doc.add_paragraph("4. Which evaluator is more accurate in terms of Medical-Images-Description?\n\n")
        doc.add_paragraph("5. What is your feedback?\n\n")

        case_output_dir = os.path.join(output_dir, case)
        os.makedirs(case_output_dir, exist_ok=True)
        docx_output_path = os.path.join(case_output_dir, f"{case}_{file_prefix}_combined_o1_r1.docx")
        doc.save(docx_output_path)