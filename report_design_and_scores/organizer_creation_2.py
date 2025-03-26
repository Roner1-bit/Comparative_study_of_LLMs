import re
import os
import csv
import shutil
from docx import Document
from docx.shared import Inches
from PIL import Image  # Ensure Pillow is installed (pip install Pillow)

def remove_invalid_chars(text):
    # Removes characters outside typical ASCII range (except \n, \r, \t)
    import re
    return re.sub(r'[^\x20-\x7E\n\r\t]+', '', text)

original_folder_path = r'D:\final_drive\(final)merged_images_with_labels_order_and_folders_mask_normalized'
new_root_folder = r'D:\final_drive\new_cases_folder'

# Define model encodings
model_encodings = {
    "deepseek-vl-small.txt": "Model A",
    "deepseek-vl-tiny.txt": "Model B",
    "deepseek-vl2.txt": "Model C",
    "diagnostic_prompt.txt": "Model D",
    "llama-11B.txt": "Model E",
    "llava-med-response11.txt": "Model F",
    "llava_answer.txt": "Model G",
    "nvlm-72b-response.txt": "Model H",
    "paligemma2-10b-report.txt": "Model I",
    "paligemma2-28b-report.txt": "Model J",
    "paligemma2-3b-report.txt": "Model K",
    "phi-3.5-vision-instruct-response.txt": "Model L",
    "qwen-vl-2b-response.txt": "Model M",
    "qwen-vl-72b-response.txt": "Model N",
    "qwen-vl-7b-response.txt": "Model O",
    "RHUH-0001.docx": "Model P"  # Example
}

# Titles to prepend with "### "
special_titles = {
    "Tumor Characteristics:",
    "Segmentation Analysis:",
    "Surgical Considerations:",
    "Clinical Summary:",
    "Recommendations:",
    "Prognostic Considerations:",
    "Follow-Up Plan:"
}

# Create the new root folder
os.makedirs(new_root_folder, exist_ok=True)

# Create a CSV listing the model encodings
csv_path = os.path.join(new_root_folder, "model_encodings.csv")
with open(csv_path, "w", newline="", encoding="utf-8") as csvfile:
    writer = csv.writer(csvfile)
    writer.writerow(["Model File", "Encoding"])
    for m_file, enc in model_encodings.items():
        writer.writerow([m_file, enc])

# Walk the original folder structure
for root, dirs, files in os.walk(original_folder_path):
    rel_path = os.path.relpath(root, original_folder_path)
    new_folder_path = os.path.join(new_root_folder, rel_path)
    os.makedirs(new_folder_path, exist_ok=True)

    # Copy and rotate only PNG images
    rotated_images = []
    for f in files:
        if f.lower().endswith(".png"):
            src_file = os.path.join(root, f)
            dst_file = os.path.join(new_folder_path, f)
            # Rotate image 90° to the right (clockwise) before saving
            with Image.open(src_file) as img:
                rotated_img = img.rotate(-90, expand=True)  # -90 for clockwise
                rotated_img.save(dst_file)
            rotated_images.append(f)

    # Load [folder_name].docx if present
    ref_docx_content = None
    folder_name = os.path.basename(root)
    folder_docx = os.path.join(root, folder_name + ".docx")
    if os.path.exists(folder_docx):
        ref_docx_content = Document(folder_docx)

    # Generate one Word doc per recognized model file
    for f in files:
        if f in model_encodings:
            ai_report_file = os.path.join(root, f)
            with open(ai_report_file, "r", encoding="utf-8", errors="replace") as ar:
                raw_content = ar.read()
            ai_report_content = remove_invalid_chars(raw_content)

            doc = Document()

            # Insert rotated images from new folder
            for img in rotated_images:
                doc.add_picture(os.path.join(new_folder_path, img), width=Inches(6.0))
                doc.add_paragraph()

            # Reference
            doc.add_heading("Reference Radiology report", level=1)
            if ref_docx_content:
                for paragraph in ref_docx_content.paragraphs:
                    new_p = doc.add_paragraph("")
                    for run in paragraph.runs:
                        cleaned_text = remove_invalid_chars(run.text)
                        # Prepend "### " if the entire run is one of the special titles
                        if cleaned_text.strip() in special_titles:
                            cleaned_text = "### " + cleaned_text
                        new_run = new_p.add_run(cleaned_text)
                        # Retain formatting
                        new_run.bold = run.bold
                        new_run.italic = run.italic
                        new_run.underline = run.underline
                        new_run.font.name = run.font.name
                        new_run.font.size = run.font.size
            else:
                doc.add_paragraph("[No reference document found]")

            # AI report
            doc.add_heading("AI report", level=1)
            doc.add_paragraph(ai_report_content)

            # Divider
            doc.add_paragraph("-----------------------------------------------------")

            # Space for answers
            doc.add_paragraph("What is your feedback about the AI report?\n")
            doc.add_paragraph("Please give a conciseness score out of 10:\n")
            doc.add_paragraph("Please give a correctness score out of 10:\n")
            doc.add_paragraph("Please give a completeness score out of 10:\n")
            doc.add_paragraph("Overall score out of 10:\n\n")

            model_letter = model_encodings[f].replace("Model ", "").lower()
            doc_filename = model_letter + ".docx"
            # Save the document in the same folder as the rotated images
            doc.save(os.path.join(new_folder_path, doc_filename))